from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ai_guru.state import AgentState
import os
import io
import datetime
from ai_guru.utils.path_utils import get_resource_path


def _format_list(items, numbered=False):
    """Formats a list of strings into a professional bulleted or numbered string."""
    if not items:
        return "-"
    if isinstance(items, str):
        return items
    
    lines = []
    for i, item in enumerate(items):
        if numbered:
            lines.append(f"{i+1}. {item}")
        else:
            lines.append(f"• {item}")
    return "\n".join(lines)


def _format_kegiatan(kegiatan_data):
    """Processes granular kegiatan JSON into a structured string."""
    if not isinstance(kegiatan_data, dict):
        return str(kegiatan_data)
    
    output = []
    
    if 'pendahuluan' in kegiatan_data:
        output.append("1. PENDAHULUAN")
        output.append(_format_list(kegiatan_data['pendahuluan']))
        output.append("")
        
    if 'inti' in kegiatan_data:
        output.append("2. KEGIATAN INTI")
        output.append(_format_list(kegiatan_data['inti']))
        output.append("")
        
    if 'penutup' in kegiatan_data:
        output.append("3. PENUTUP")
        output.append(_format_list(kegiatan_data['penutup']))
        
    return "\n".join(output)


def _format_asesmen(asesmen_data):
    """Processes granular asesmen JSON into a structured string."""
    if not isinstance(asesmen_data, dict):
        return str(asesmen_data)
    
    output = []
    if 'formatif' in asesmen_data:
        output.append("Asesmen Formatif:")
        output.append(_format_list(asesmen_data['formatif']))
        output.append("")
    
    if 'sumatif' in asesmen_data:
        output.append("Asesmen Sumatif:")
        output.append(_format_list(asesmen_data['sumatif']))
        
    return "\n".join(output)


def _build_soal_doc_in_memory(state: AgentState, context: dict) -> bytes:
    """
    Always generate Soal as DOCX from scratch using python-docx.
    This is the most reliable approach and avoids Jinja/template issues.
    """
    doc = Document()
    doc.add_heading(f"BANK SOAL: {state['topic']}", 0)
    
    meta = doc.add_paragraph()
    meta.add_run(f"Mata Pelajaran: {state['subject']}  |  Kelas: {state['grade_level']}").bold = True
    doc.add_paragraph("─" * 60)

    questions = context.get('questions', [])
    
    for q in questions:
        q_num = q.get('id', '?')
        q_text = q.get('question', '')
        q_type = q.get('type', '')
        options = q.get('options') or []
        
        # Question paragraph
        p = doc.add_paragraph()
        p.add_run(f"{q_num}. [{q_type}] ").bold = True
        p.add_run(q_text)
        
        # Options
        for opt in options:
            if opt:
                doc.add_paragraph(f"   {opt}", style='List Bullet')
        
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading("KUNCI JAWABAN", 1)
    
    for q in questions:
        q_num = q.get('id', '?')
        answer = q.get('answer_key', '-')
        taxonomy = q.get('taxonomy', '-')
        p = doc.add_paragraph()
        p.add_run(f"{q_num}. ").bold = True
        p.add_run(f"{answer} ")
        p.add_run(f"(Taksonomi: {taxonomy})").italic = True
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_rpp_doc_in_memory(state: AgentState, context: dict) -> bytes:
    """
    Always generate RPP as DOCX from scratch using python-docx.
    """
    doc = Document()
    doc.add_heading(f"MODUL AJAR / RPP", 0)
    doc.add_heading(f"{state['topic']}", 1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Shading Accent 1'
    cells = info_table.rows
    cells[0].cells[0].text = 'Mata Pelajaran'
    cells[0].cells[1].text = str(state['subject'])
    cells[1].cells[0].text = 'Jenjang/Kelas'
    cells[1].cells[1].text = str(state['grade_level'])
    cells[2].cells[0].text = 'Nama Guru'
    cells[2].cells[1].text = context.get('guru', '-')
    cells[3].cells[0].text = 'Sekolah'
    cells[3].cells[1].text = context.get('sekolah', '-')
    cells[4].cells[0].text = 'Tahun Ajaran'
    cells[4].cells[1].text = context.get('tahun', '-')
    
    doc.add_heading("A. Tujuan Pembelajaran", 2)
    tujuan = context.get('rpp_tujuan', '-')
    doc.add_paragraph(tujuan)
    
    doc.add_heading("B. Langkah Kegiatan", 2)
    kegiatan = context.get('rpp_kegiatan', '-')
    doc.add_paragraph(kegiatan)
    
    doc.add_heading("C. Asesmen", 2)
    asesmen = context.get('rpp_asesmen', '-')
    doc.add_paragraph(asesmen)

    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def format_document(state: AgentState) -> AgentState:
    """
    Node to format the RPP and Questions into DOCX files.
    Saves file to disk AND stores bytes in state for direct download.
    """
    print(f"Formatting documents for: {state['topic']}")
    today = datetime.date.today().strftime("%d %B %Y")

    
    admin_data = state.get('admin_data') or {}
    
    # Prepare Context for Templating
    context = {
        'topic': state['topic'],
        'subject': state.get('subject', '-'),
        'grade': state['grade_level'],
        'sekolah': admin_data.get('sekolah', '-'),
        'guru': admin_data.get('guru', '-'),
        'nip': admin_data.get('nip', '-'),
        'kepsek': admin_data.get('kepsek', '-'),
        'nip_kepsek': '................',
        'tahun': admin_data.get('tahun_ajar', '-'),
        'date': today,
        'rpp_tujuan': "",
        'rpp_kegiatan': "",
        'rpp_asesmen': "",
        'questions': []
    }
    
    # Sanitize questions to prevent Jinja/NoneType errors
    if state.get('questions'):
        safe_questions = []
        for q in state['questions']:
            if not isinstance(q, dict):
                continue
            q_safe = dict(q)
            options = q_safe.get('options')
            if options is None:
                q_safe['options'] = []
            elif isinstance(options, str):
                q_safe['options'] = [options]
            elif not isinstance(options, list):
                q_safe['options'] = list(options)
            safe_questions.append(q_safe)
        context['questions'] = safe_questions
    
    if state.get('rpp'):
        rpp_raw = state['rpp']
        context['rpp_tujuan'] = _format_list(rpp_raw.get('tujuan_pembelajaran', []), numbered=True)
        context['rpp_kegiatan'] = _format_kegiatan(rpp_raw.get('langkah_kegiatan', {}))
        context['rpp_asesmen'] = _format_asesmen(rpp_raw.get('asesmen', {}))


    base_dir = os.path.dirname(os.path.abspath(__file__))
    # Go up to project root (src/agents -> src -> root)
    project_root = os.path.abspath(os.path.join(base_dir, '..', '..'))

    # --- 1. RPP Generation ---
    if state.get('rpp'):
        try:
            template_rpp_path = get_resource_path('templates/template_rpp.docx')
            if template_rpp_path.exists():
                print(f"Using RPP Template: {template_rpp_path}")
                doc = DocxTemplate(str(template_rpp_path))
                doc.render(context)
                buf = io.BytesIO()
                doc.save(buf)
                state['rpp_docx'] = buf.getvalue()
            else:
                print(f"RPP Template not found at {template_rpp_path}, using fallback.")
                rpp_bytes = _build_rpp_doc_in_memory(state, context)
                state['rpp_docx'] = rpp_bytes
        except Exception as e:
            print(f"Error generating RPP: {e}")
            state['logs'].append(f"RPP Formatting Error: {str(e)}")

    # --- 2. Soal Generation ---
    if state.get('questions'):
        try:
            template_soal_path = get_resource_path('templates/template_soal.docx')
            if template_soal_path.exists():
                print(f"Using Soal Template: {template_soal_path}")
                doc = DocxTemplate(str(template_soal_path))
                doc.render(context)
                buf = io.BytesIO()
                doc.save(buf)
                state['soal_docx'] = buf.getvalue()
            else:
                print(f"Soal Template not found at {template_soal_path}, using fallback.")
                soal_bytes = _build_soal_doc_in_memory(state, context)
                state['soal_docx'] = soal_bytes
        except Exception as e:
            print(f"Soal Template failed: {e}. Falling back to programmatic generation.")
            try:
                soal_bytes = _build_soal_doc_in_memory(state, context)
                state['soal_docx'] = soal_bytes
            except Exception as e2:
                print(f"Fatal error generating Soal: {e2}")
                state['logs'].append(f"Soal Formatting Error: {str(e2)}")

        
    state['logs'].append("Documents generated.")
    return state
