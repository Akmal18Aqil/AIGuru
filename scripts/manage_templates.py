from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_instructions_doc(filename):
    doc = Document()
    doc.add_heading('PANDUAN EDIT TEMPLATE', 0)
    doc.add_paragraph("Anda bisa mengedit file 'template_rpp.docx' dan 'template_soal.docx' sesuai keinginan.")
    doc.add_paragraph("Jangan hapus teks dengan tanda kurung kurawal ganda, contoh: {{ guru }}.")
    doc.add_paragraph("Daftar Placeholder:")
    
    items = [
        "{{ topic }} - Topik Materi",
        "{{ subject }} - Mata Pelajaran",
        "{{ grade }} - Jenjang Kelas",
        "{{ sekolah }} - Nama Sekolah",
        "{{ guru }} - Nama Guru",
        "{{ nip }} - NIP Guru",
        "{{ kepsek }} - Nama Kepala Sekolah",
        "{{ tahun }} - Tahun Ajaran",
        "{{ rpp_tujuan }} - List Tujuan Pembelajaran",
        "{{ rpp_kegiatan }} - Deskripsi Kegiatan",
        "{{ rpp_asesmen }} - Deskripsi Asesmen",
        "{{ questions }} - List Soal (Looping otomatis)"
    ]
    
    for item in items:
        doc.add_paragraph(item, style='List Bullet')
        
    doc.save(filename)

def create_rpp_template(filename):
    doc = Document()
    
    # Header KOP (User can replace this)
    header = doc.add_heading('{{ sekolah }}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Alamat: Jalan Pendidikan No. 1 Indonesia')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("-----------------------------------------------------------------------------------------------------------------------")
    
    # Title
    title = doc.add_heading('MODUL AJAR / RPP', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Info Table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    table.cell(0,0).text = "Mata Pelajaran"
    table.cell(0,1).text = ":"
    table.cell(0,2).text = "{{ subject }}"
    
    table.cell(1,0).text = "Kelas / Fase"
    table.cell(1,1).text = ":"
    table.cell(1,2).text = "{{ grade }}"
    
    table.cell(2,0).text = "Topik"
    table.cell(2,1).text = ":"
    table.cell(2,2).text = "{{ topic }}"
    
    table.cell(3,0).text = "Tahun Ajaran"
    table.cell(3,1).text = ":"
    table.cell(3,2).text = "{{ tahun }}"
    
    doc.add_paragraph()
    
    # Content
    doc.add_heading('A. Tujuan Pembelajaran', level=2)
    doc.add_paragraph("{{ rpp_tujuan }}") # In docxtpl, rich text insertion is handled via context
    
    doc.add_heading('B. Kegiatan Pembelajaran', level=2)
    doc.add_paragraph("{{ rpp_kegiatan }}")
    
    doc.add_heading('C. Asesmen', level=2)
    doc.add_paragraph("{{ rpp_asesmen }}")
    
    # Signature
    doc.add_paragraph()
    table_sig = doc.add_table(rows=3, cols=2)
    table_sig.autofit = True
    
    # Date
    table_sig.cell(0,1).text = "..........., {{ date }}"
    
    table_sig.cell(0,0).text = "Mengetahui,\nKepala Sekolah"
    table_sig.cell(0,1).text += "\nGuru Mata Pelajaran"
    
    table_sig.cell(1,0).text = "\n\n\n"
    
    table_sig.cell(2,0).text = "{{ kepsek }}\nNIP. {{ nip_kepsek }}"
    table_sig.cell(2,1).text = "{{ guru }}\nNIP. {{ nip }}"
    
    doc.save(filename)

def create_soal_template(filename):
    doc = Document()
    
    doc.add_heading('BANK SOAL {{ sekolah }}', 0)
    doc.add_paragraph('Mapel: {{ subject }} | Kelas: {{ grade }}')
    doc.add_heading('{{ topic }}', 1)
    doc.add_paragraph("--------------------------------------------------")
    
    # Loop for questions
    # docxtpl uses jinja syntax inside word
    # We will simulate the jinja tag here as text
    p = doc.add_paragraph("{% for q in questions %}")
    
    # Question Body
    doc.add_paragraph("{{ q.id }}. {{ q.question }}")
    
    # Options loop
    doc.add_paragraph("{% for opt in q.options %}")
    doc.add_paragraph("- {{ opt }}")
    doc.add_paragraph("{% endfor %}")
    
    doc.add_paragraph() # Spacer
    doc.add_paragraph("{% endfor %}")
    
    # Keys
    doc.add_page_break()
    doc.add_heading('KUNCI JAWABAN', 1)
    doc.add_paragraph("{% for q in questions %}")
    doc.add_paragraph("{{ q.id }}. {{ q.answer_key }} (Taksonomi: {{ q.taxonomy }})")
    doc.add_paragraph("{% endfor %}")

    doc.save(filename)

if __name__ == "__main__":
    os.makedirs("templates", exist_ok=True)
    create_rpp_template("templates/template_rpp.docx")
    create_soal_template("templates/template_soal.docx")
    create_instructions_doc("templates/BACA_PANDUAN_TEMPLATE.docx")
    print("Templates created in /templates folder.")
