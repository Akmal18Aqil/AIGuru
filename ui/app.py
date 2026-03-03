import streamlit as st
import pandas as pd
import sys
import os

# Add src to pythonpath so imports work
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from ai_guru.main_graph import app_graph
from ai_guru.utils.licensing import LicenseManager
from ai_guru.utils.document_loader import load_document_text
from ai_guru.config.api_key_manager import APIKeyManager

st.set_page_config(page_title="SiGURU (AI Assistant)", layout="wide", page_icon="🎓")

# Absolute path to project root (ui/app.py is inside ui/ subfolder)
PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Initialize API Key Manager
api_manager = APIKeyManager()

# Check if setup is completed - if not, redirect to setup page
if not api_manager.is_setup_completed():
    st.warning("⚙️ Aplikasi belum dikonfigurasi. Mengalihkan ke Setup Wizard...")
    st.info("Silakan tunggu beberapa detik...")
    st.switch_page("pages/0_⚙️_Setup.py")


# === HELPER FUNCTIONS FOR SMART CONFLICT FIX ===

def render_smart_fix(conflict, guru, kelas_to_move, idx):
    """Render smart auto-fix options with validation"""
    from ai_guru.utils.scheduler_logic import find_available_slots
    from ai_guru.utils.conflict_detector import detect_hard_conflicts
    
    df_current = pd.DataFrame(st.session_state['main_jadwal_result'])
    free_slots = find_available_slots(df_current, guru, kelas_to_move)
    
    if free_slots:
        st.success(f"✅ **Ditemukan {len(free_slots)} slot alternatif untuk {kelas_to_move}**")
        
        # Show top 3 options
        cols = st.columns(min(3, len(free_slots)))
        for i, (hari, jam) in enumerate(free_slots[:3]):
            with cols[i]:
                if st.button(
                    f"📍 {hari} Jam {jam}",
                    key=f"fix_{idx}_{guru}_{kelas_to_move}_{hari}_{jam}",
                    use_container_width=True
                ):
                    # Execute fix with re-validation
                    success = execute_fix_and_revalidate(
                        conflict=conflict,
                        new_hari=hari,
                        new_jam=jam,
                        guru=guru,
                        kelas_to_move=kelas_to_move
                    )
                    
                    if success:
                        st.success("✅ Berhasil diperbaiki!")
                        st.rerun()
                    else:
                        st.error("❌ Gagal! Slot masih bentrok.")
    else:
        st.error(f"❌ Tidak ada slot kosong untuk **{kelas_to_move}**")
        st.info("💡 Saran: Kurangi beban mengajar atau tambah hari kerja")


def execute_fix_and_revalidate(conflict, new_hari, new_jam, guru, kelas_to_move):
    """Execute fix and re-validate to ensure no new conflicts"""
    from ai_guru.utils.conflict_detector import detect_hard_conflicts
    
    # 1. Create updated schedule
    updated_jadwal = []
    found_entry = False
    
    for entry in st.session_state['main_jadwal_result']:
        entry_copy = entry.copy()
        
        # Find the conflicting entry to move
        if (entry_copy['guru'] == guru and 
            entry_copy['kelas'] == kelas_to_move and
            entry_copy['hari'] == conflict['hari'] and
            entry_copy['jam_ke'] == conflict['jam_ke']):
            
            entry_copy['hari'] = new_hari
            entry_copy['jam_ke'] = new_jam
            found_entry = True
        
        updated_jadwal.append(entry_copy)
    
    if not found_entry:
        return False
    
    # 2. RE-VALIDATE (CRITICAL!)
    new_conflicts = detect_hard_conflicts(updated_jadwal)
    
    if not new_conflicts['has_conflict']:
        # SUCCESS! Apply changes
        st.session_state['main_jadwal_result'] = updated_jadwal
        
        # Re-run conflict detection to update UI
        st.session_state['main_jadwal_conflicts'] = new_conflicts
        return True
    else:
        # FAILED: New slot also has conflict
        return False


# --- CUSTOM CSS ---
st.markdown("""
<style>
    /* Hide default Streamlit sidebar navigation */
    [data-testid="stSidebarNav"] {
        display: none;
    }
    
    /* Premium Sidebar Styling */
    section[data-testid="stSidebar"] {
        background-color: #0e1117;
    }
    
    /* Button Base Style */
    .stButton > button {
        width: 100%;
        border-radius: 12px;
        height: 3.5em;
        font-weight: 600;
        border: 1px solid rgba(255, 255, 255, 0.1);
        background-color: rgba(255, 255, 255, 0.05);
        color: white;
        transition: all 0.3s ease;
        text-align: left;
        padding-left: 1.5em;
        margin-bottom: 0.5em;
    }
    
    /* Button Hover Effect */
    .stButton > button:hover {
        border-color: #4A90E2;
        background-color: rgba(74, 144, 226, 0.1);
        transform: translateY(-2px);
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.2);
    }
    
    /* Active Button Style (Primary Type) */
    .stButton > button[kind="primary"] {
        background: linear-gradient(90deg, #4A90E2 0%, #357ABD 100%);
        border: none;
        box-shadow: 0 4px 15px rgba(74, 144, 226, 0.3);
        color: white !important;
    }
    
    .stButton > button[kind="primary"]:hover {
        background: linear-gradient(90deg, #357ABD 0%, #4A90E2 100%);
        transform: translateY(-2px);
    }

    /* Clean divider */
    .stDivider {
        margin-top: 1.5rem !important;
        margin-bottom: 1.5rem !important;
    }
    
    .big-font {
        font-size: 20px !important;
    }
</style>
""", unsafe_allow_html=True)

# --- STATE MANAGEMENT ---
if 'authenticated' not in st.session_state:
    st.session_state['authenticated'] = False
if 'current_page' not in st.session_state:
    st.session_state['current_page'] = "Home"

# --- LICENSE CHECK (NEW!) ---
# This happens AFTER setup check but BEFORE user authentication
license_status = api_manager.get_license_status()

if not license_status['has_license']:
    st.error("❌ **License Key Tidak Ditemukan!**")
    st.warning("Aplikasi ini membutuhkan license key yang valid untuk berfungsi.")
    st.info("Silakan jalankan Setup Wizard untuk memasukkan license key.")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        if st.button("⚙️ Buka Setup Wizard", type="primary"):
            st.switch_page("pages/0_⚙️_Setup.py")
    
    st.markdown("---")
    st.markdown("""
    ### ℹ️ Tentang License Key
    
    License key diperlukan untuk:
    - ✅ Memvalidasi penggunaan resmi SiGURU
    - 📊 Akses ke semua fitur
    - 🔄 Update dan support
    
    **Cara mendapatkan license:**
    - 📧 Email: sales@siguru.app
    - 💬 WhatsApp: +62-XXX-XXXX-XXXX
    """)
    st.stop()

if not license_status['is_valid']:
    st.error("❌ **License Key Tidak Valid atau Sudah Expired!**")
    st.warning(f"License Anda: `{license_status['license_key']}`")
    
    st.info("""
    **Kemungkinan penyebab:**
    - License sudah expired
    - License tidak terdaftar di sistem
    - Koneksi internet terputus (jika menggunakan validasi online)
    
    **Solusi:**
    - Hubungi support untuk perpanjangan license
    - Atau jalankan setup ulang untuk memasukkan license key baru
    """)
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📧 Hubungi Support"):
            st.info("Email: support@siguru.app")
    with col2:
        if st.button("⚙️ Setup Ulang"):
            # Reset setup to re-enter license
            if st.button("Konfirmasi Reset Setup?", type="secondary"):
                api_manager.reset_setup()
                st.switch_page("pages/0_⚙️_Setup.py")
    
    st.stop()

# --- LICENSING CHECK ---
def check_login():
    key = st.session_state.get('license_input', '')
    manager = LicenseManager()
    if manager.verify_license(key):
        st.session_state['authenticated'] = True
    else:
        st.error("Lisensi tidak valid atau tidak aktif.")

# Auto-authenticate if license is already valid in config
if license_status['is_valid'] and not st.session_state['authenticated']:
    st.session_state['authenticated'] = True

if not st.session_state['authenticated']:
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        st.title("🔐 SiGURU Activation")
        st.markdown("---")
        st.markdown("Masukkan Kode Lisensi Anda untuk masuk.")
        st.text_input("License Key", key="license_input", type="password")
        st.button("Masuk", on_click=check_login)
    st.stop()

# --- NAVIGATION ---
with st.sidebar:
    st.markdown("""
    <div style="text-align: center; padding: 1em 0;">
        <h1 style="color: #4A90E2; margin-bottom: 0;">🎓 SiGURU</h1>
        <p style="color: #888; font-size: 0.8em; margin-top: 0;">Smart AI Assistant for Teachers</p>
    </div>
    """, unsafe_allow_html=True)
    
    # API Key Status Indicator (More compact)
    api_status = api_manager.get_status()
    if api_status['active']:
        st.caption(f"🟢 API {api_status['type']} Active")
    else:
        st.error("⚠️ API Key inactive")
        if st.button("⚙️ Setup Ulang"):
            st.switch_page("pages/0_⚙️_Setup.py")
    
    st.divider()
    
    # Navigation Buttons with Active State
    if st.button("🏠 Beranda", 
                 type="primary" if st.session_state['current_page'] == "Home" else "secondary",
                 use_container_width=True):
        st.session_state['current_page'] = "Home"
        st.rerun()
        
    if st.button("📚 Generator RPP", 
                 type="primary" if st.session_state['current_page'] == "Modul Ajar" else "secondary",
                 use_container_width=True):
        st.session_state['current_page'] = "Modul Ajar"
        st.rerun()
        
    if st.button("📝 Generator Soal", 
                 type="primary" if st.session_state['current_page'] == "Generator Soal" else "secondary",
                 use_container_width=True):
        st.session_state['current_page'] = "Generator Soal"
        st.rerun()
        
    if st.button("📅 Jadwal Pelajaran", 
                 type="primary" if st.session_state['current_page'] == "Jadwal" else "secondary",
                 use_container_width=True):
        st.session_state['current_page'] = "Jadwal"
        st.rerun()
        
    st.divider()
    
    # Footer Section
    with st.expander("⚙️ Pengaturan & Info"):
        if st.button("🔧 Setup Wizard"):
            st.switch_page("pages/0_⚙️_Setup.py")
        
        # Troubleshooting Log Viewer
        if st.checkbox("🔍 Tampilkan Debug Log"):
            log_file = Path(__file__).parent.parent / "logs" / "siguru.log"
            if log_file.exists():
                with open(log_file, "r", encoding='utf-8') as f:
                    # Show last 20 lines
                    lines = f.readlines()
                    st.code("".join(lines[-20:]))
            else:
                st.caption("Log file not found yet.")
                    
        st.caption("Versi: 1.1.0 (Prod-Ready)")
        st.caption("© 2024 SiGURU AI")

# --- PAGE: HOME ---
if st.session_state['current_page'] == "Home":
    st.title("Selamat Datang di SiGURU")
    st.markdown("### Asisten Cerdas untuk Guru Hebat")
    
    col1, col2 = st.columns(2)
    with col1:
        st.info("### 📚 RPP Builder\nBuat RPP Otomatis Kurikulum Merdeka.")
        if st.button("Buka RPP Builder"):
            st.session_state['current_page'] = "Modul Ajar"
            st.rerun()
        st.info("### 📝 Generator Soal\nBuat / Remix Soal Ujian dengan AI.")
        if st.button("Buka Generator Soal"):
            st.session_state['current_page'] = "Generator Soal"
            st.rerun()
            
    with col2:
        st.success("### 📅 Smart Scheduler\nSusun jadwal pelajaran anti-bentrok untuk seluruh sekolah.")
        if st.button("Buka Scheduler"):
            st.session_state['current_page'] = "Jadwal"
            st.rerun()

# --- PAGE: MODUL AJAR (RPP ONLY) ---
elif st.session_state['current_page'] == "Modul Ajar":
    st.title("📚 Generator Modul Ajar (RPP)")
    
    col_input, col_preview = st.columns([1, 1])
    
    with col_input:
        st.subheader("Konfigurasi")
        topic = st.text_input("Topik / Materi", "Sistem Pencernaan Manusia")
        grade_level = st.selectbox("Jenjang", ["SD Kelas 1-6", "SMP Kelas 7-9", "SMA Kelas 10-12"])
        subject = st.text_input("Mata Pelajaran", "Ilmu Pengetahuan Alam (IPA)")
        
        with st.expander("Identitas Sekolah", expanded=True):
            guru = st.text_input("Nama Guru", "Siti Aminah, S.Pd", key="guru_rpp")
            nip = st.text_input("NIP", "19800101...", key="nip_rpp")
            sekolah = st.text_input("Nama Sekolah", "SDN 01 Jakarta", key="sekolah_rpp")
            kepsek = st.text_input("Kepala Sekolah", "Budi Santoso, M.Pd", key="kepsek_rpp")
            tahun_ajar = st.text_input("Tahun Ajaran", "2024/2025", key="ta_rpp")
            
        generate_btn = st.button("🚀 Generate RPP")

    with col_preview:
        st.subheader("Hasil & Preview")
        if generate_btn:
            api_key = api_manager.get_api_key()
            if not api_key:
                st.error("API Key tidak ditemukan! Silakan ke halaman Setup.")
                if st.button("⚙️ Buka Setup", key="setup_rpp"):
                    st.switch_page("pages/0_⚙️_Setup.py")
            else:
                os.environ["GOOGLE_API_KEY"] = api_key
                with st.spinner("Sedang menyusun RPP..."):
                    initial_state = {
                        "topic": topic, "grade_level": grade_level, "subject": subject,
                        "admin_data": {"guru": guru, "nip": nip, "sekolah": sekolah, "kepsek": kepsek, "tahun_ajar": tahun_ajar},
                        "source_text": "", "use_rag": False, "generation_mode": "rpp_only",
                        "rpp": None, "questions": [], "status": "Running", "logs": []
                    }
                    
                    st.success("RPP Selesai Disusun!")
                    
                    # === LOGS & STATUS ===
                    with st.expander("🎓 Detail Proses & Log AI", expanded=False):
                        if 'logs' in final_state:
                            for log in final_state['logs']:
                                if any(x in log for x in ["Error", "Failed", "Critical"]):
                                    st.error(f"❌ {log}")
                                elif "Warning" in log:
                                    st.warning(f"⚠️ {log}")
                                else:
                                    st.info(f"🔹 {log}")
                    
                    # Build DOCX in-memory
                    if final_state.get('rpp'):
                        import io
                        from docx import Document
                        rpp = final_state['rpp']
                        doc = Document()
                        doc.add_heading(f"Modul Ajar / RPP: {topic}", 0)
                        doc.add_heading(subject, 1)
                        
                        # Identity Table
                        tbl = doc.add_table(rows=5, cols=2)
                        tbl.style = 'Light Shading Accent 1'
                        rows = tbl.rows
                        rows[0].cells[0].text = 'Mata Pelajaran';  rows[0].cells[1].text = subject
                        rows[1].cells[0].text = 'Jenjang';         rows[1].cells[1].text = grade_level
                        rows[2].cells[0].text = 'Guru';            rows[2].cells[1].text = guru
                        rows[3].cells[0].text = 'Sekolah';         rows[3].cells[1].text = sekolah
                        rows[4].cells[0].text = 'Tahun Ajaran';    rows[4].cells[1].text = tahun_ajar
                        doc.add_paragraph()
                        
                        doc.add_heading('A. Tujuan Pembelajaran', 2)
                        tujuan = rpp.get('tujuan_pembelajaran', [])
                        if isinstance(tujuan, list):
                            for tp in tujuan: doc.add_paragraph(tp, style='List Bullet')
                        else:
                            doc.add_paragraph(str(tujuan))
                        
                        doc.add_heading('B. Langkah Kegiatan', 2)
                        doc.add_paragraph(rpp.get('langkah_kegiatan', '-'))
                        
                        doc.add_heading('C. Asesmen', 2)
                        doc.add_paragraph(rpp.get('asesmen', '-'))
                        
                        buf_rpp = io.BytesIO()
                        doc.save(buf_rpp)
                        safe_topic = topic.replace(' ', '_').replace('/', '-')
                        st.download_button(
                            label="📄 Download RPP (DOCX)",
                            data=buf_rpp.getvalue(),
                            file_name=f"RPP_{safe_topic}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        st.info(f"Tujuan Pembelajaran: {tujuan}")

# --- PAGE: GENERATOR SOAL ---
elif st.session_state['current_page'] == "Generator Soal":
    st.title("📝 Generator Soal Ujian")
    
    col_input, col_preview = st.columns([1, 1])
    
    with col_input:
        st.subheader("Sumber Materi")
        topic = st.text_input("Topik Referensi", "Sistem Pencernaan Manusia")
        grade_level = st.selectbox("Jenjang Akademik", ["SD Kelas 1-6", "SMP Kelas 7-9", "SMA Kelas 10-12"])
        
        with st.expander("Modifikasi dari File Lama (Opsional)"):
            st.info("Gunakan ini jika Anda ingin membuat variasi soal dari dokumen lama.")
            uploaded_file = st.file_uploader("Upload Soal/Materi (PDF/DOCX)", type=['pdf', 'docx'])
            use_rag = st.checkbox("Remix soal dari file ini", value=False, disabled=not uploaded_file)
            
        num_questions = st.number_input("Jumlah Soal", min_value=1, max_value=100, value=50)
        
        qt_options = ["Pilihan Ganda", "Isian Singkat", "Uraian", "HOTS (Higher Order Thinking Skills)", "Studi Kasus", "Benar/Salah", "Menjodohkan"]
        selected_types = st.multiselect(
            "Jenis Soal",
            options=qt_options,
            default=["Pilihan Ganda", "Isian Singkat", "Uraian"],
            help="Pilih satu atau lebih jenis soal yang ingin dibuat."
        )
            
        generate_soal_btn = st.button(f"🚀 Generate {num_questions} Soal")

    with col_preview:
        st.subheader("Soal Tersedia")
        if generate_soal_btn:
            api_key = api_manager.get_api_key()
            if not api_key:
                st.error("API Key tidak ditemukan! Silakan ke halaman Setup.")
                if st.button("⚙️ Buka Setup", key="setup_soal"):
                    st.switch_page("pages/0_⚙️_Setup.py")
            else:
                os.environ["GOOGLE_API_KEY"] = api_key
                with st.spinner(f"Merumuskan {num_questions} soal... (Bisa butuh 1-2 menit)"):
                    source_text = ""
                    if uploaded_file and use_rag:
                        with open("temp_upload", "wb") as f: f.write(uploaded_file.getbuffer())
                        source_text = load_document_text("temp_upload")
                        
                    initial_state = {
                        "topic": topic, "grade_level": grade_level, "subject": "-",
                        "admin_data": {},
                        "source_text": source_text, "use_rag": use_rag, "generation_mode": "soal_only",
                        "num_questions": num_questions,
                        "question_types": selected_types if selected_types else ["Pilihan Ganda"],
                        "rpp": None, "questions": [], "status": "Running", "logs": []
                    }
                    
                    final_state = app_graph.invoke(initial_state)
                    st.success("Selesai Generasi Soal!")
                    
                    # === LOGS & STATUS ===
                    with st.expander("🎓 Detail Proses & Log AI", expanded=False):
                        if 'logs' in final_state:
                            for log in final_state['logs']:
                                if any(x in log for x in ["Error", "Failed", "Critical"]):
                                    st.error(f"❌ {log}")
                                elif "Warning" in log:
                                    st.warning(f"⚠️ {log}")
                                else:
                                    st.info(f"🔹 {log}")

                    # Build DOCX in-memory (same approach as Jadwal)
                    questions = final_state.get('questions', [])
                    if questions:
                        import io
                        from docx import Document
                        doc = Document()
                        doc.add_heading(f"Bank Soal: {topic}", 0)
                        p = doc.add_paragraph()
                        p.add_run(f"Kelas: {grade_level}  |  Jumlah Soal: {len(questions)}").bold = True
                        doc.add_paragraph("─" * 55)
                        
                        for q in questions:
                            if not isinstance(q, dict): continue
                            q_num = q.get('id', '?')
                            q_text = q.get('question', '')
                            q_type = q.get('type', '')
                            options = q.get('options') or []
                            
                            p = doc.add_paragraph()
                            p.add_run(f"{q_num}. [{q_type}] ").bold = True
                            p.add_run(q_text)
                            
                            for opt in options:
                                if opt:
                                    doc.add_paragraph(f"   {opt}", style='List Bullet')
                            doc.add_paragraph()
                        
                        doc.add_page_break()
                        doc.add_heading("Kunci Jawaban", 1)
                        for q in questions:
                            if not isinstance(q, dict): continue
                            p = doc.add_paragraph()
                            p.add_run(f"{q.get('id', '?')}. ").bold = True
                            p.add_run(f"{q.get('answer_key', '-')} ")
                            p.add_run(f"(Taksonomi: {q.get('taxonomy', '-')})").italic = True
                        
                        buf_soal = io.BytesIO()
                        doc.save(buf_soal)
                        safe_topic = topic.replace(' ', '_').replace('/', '-')
                        st.download_button(
                            label="📝 Download Soal (DOCX)",
                            data=buf_soal.getvalue(),
                            file_name=f"Soal_{safe_topic}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    else:
                        st.warning("Soal tidak berhasil dibuat. Coba lagi.")

# --- PAGE: JADWAL PELAJARAN (New & Improved) ---
elif st.session_state['current_page'] == "Jadwal":
    st.title("📅 Smart Scheduler School-Wide")
    
    tabs = st.tabs(["1. Data Guru", "2. Data Kelas", "3. Generate & Hasil"])
    
    with tabs[0]:
        st.info("Masukkan daftar guru, mata pelajaran yang diampu, dan beban jam.")
        
        if 'df_guru' not in st.session_state:
            st.session_state['df_guru'] = pd.DataFrame(
                [
                    {"Nama": "Pak Ahmad", "Mapel": "Matematika", "Jam/Minggu": 24},
                    {"Nama": "Bu Siti", "Mapel": "IPA", "Jam/Minggu": 20},
                    {"Nama": "Pak Budi", "Mapel": "B. Indonesia", "Jam/Minggu": 22},
                    {"Nama": "Bu Rina", "Mapel": "B. Inggris", "Jam/Minggu": 20},
                    {"Nama": "Pak Joko", "Mapel": "IPS", "Jam/Minggu": 18},
                ]
            )
            
        st.session_state['df_guru'] = st.data_editor(
            st.session_state['df_guru'], 
            num_rows="dynamic", 
            use_container_width=True
        )
        
    with tabs[1]:
        st.info("Daftar Kelas dan kuota jam per minggu.")
        if 'df_kelas' not in st.session_state:
            st.session_state['df_kelas'] = pd.DataFrame(
                [
                    {"Kelas": "VII-A", "Wali Kelas": "Pak Ahmad"},
                    {"Kelas": "VII-B", "Wali Kelas": "Bu Siti"},
                    {"Kelas": "VIII-A", "Wali Kelas": "Pak Budi"},
                    {"Kelas": "VIII-B", "Wali Kelas": "Bu Rina"},
                    {"Kelas": "IX-A", "Wali Kelas": "Pak Joko"},
                    {"Kelas": "IX-B", "Wali Kelas": "-"},
                ]
            )
        st.session_state['df_kelas'] = st.data_editor(
            st.session_state['df_kelas'], 
            num_rows="dynamic", 
            use_container_width=True
        )

    with tabs[2]:
        st.subheader("Eksekusi Jadwal & Aturan")

        col_config, col_action = st.columns([2, 1])
        
        with col_config:
            default_slots = (
                "Jam 1: 07:00-07:45\n"
                "Jam 2: 07:45-08:30\n"
                "Jam 3: 08:30-09:15\n"
                "Istirahat: 09:15-09:30\n"
                "Jam 4: 09:30-10:15"
            )
            time_slots = st.text_area(
                "Slot Waktu (Jam Pelajaran)", 
                default_slots,
                height=150,
                help="Tentukan slot waktu harian."
            )
            
            default_constraints = (
                "Pak Joko | tidak_hadir: Rabu\n"
                "Bu Siti | hanya_hadir: Rabu, Kamis"
            )
            constraints = st.text_area(
                "Constraint / Aturan Khusus", 
                default_constraints,
                help="Instruksi khusus: (tidak_hadir, hanya_hadir, dll)."
            )

        with col_action:
            st.write("### Siap Generate?")
            run_btn = st.button("⚡ Generate Jadwal Otomatis", type="primary")
        
        if run_btn:
            # Get API key from manager
            api_key = api_manager.get_api_key()
            if not api_key:
                st.error("API Key tidak ditemukan! Silakan ke halaman Setup.")
                if st.button("⚙️ Buka Setup"):
                    st.switch_page("pages/0_⚙️_Setup.py")
            else:
                os.environ["GOOGLE_API_KEY"] = api_key
                # Convert DataFrames to JSON-compatible lists
                teachers = st.session_state['df_guru'].to_dict('records')
                classes = st.session_state['df_kelas'].to_dict('records')
                
                with st.spinner("AI sedang menyusun jadwal... (Bisa 1-3 menit)"):
                    # State for Jadwal
                    sched_state = {
                        "topic": "Jadwal",
                        "grade_level": "-",
                        "subject": "-",
                        "admin_data": {},
                        "use_rag": False,
                        "source_text": "",
                        "generation_mode": "jadwal",
                        "rpp": None,
                        "questions": [],
                        "jadwal_mode": True,
                        "jadwal_teachers": teachers,
                        "jadwal_classes": classes,
                        "jadwal_time_slots": time_slots,
                        "jadwal_constraints": constraints,
                        "logs": []
                    }
                    
                    try:
                        result = app_graph.invoke(sched_state)
                        
                        # Store in SESSION STATE
                        if result.get('jadwal_result'):
                            st.session_state['main_jadwal_result'] = result['jadwal_result']
                        if result.get('jadwal_conflicts'):
                            st.session_state['main_jadwal_conflicts'] = result['jadwal_conflicts']
                            
                        st.success("Jadwal Berhasil Disusun!")
                        st.rerun()
                    
                    except Exception as e:
                        st.error(f"Terjadi kesalahan: {e}")

        # --- DISPLAY RESULTS (Persisted) ---
        if 'main_jadwal_result' in st.session_state:
            st.divider()
            st.subheader("📅 Hasil Jadwal")
            
            # Convert to DataFrame
            df_jadwal = pd.DataFrame(st.session_state['main_jadwal_result'])
            st.dataframe(df_jadwal, use_container_width=True)
            
            # Excel Export
            import io
            buffer = io.BytesIO()
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                # Sheet 1: All Data
                df_jadwal.to_excel(writer, index=False, sheet_name='Semua Jadwal')
                
                # Sheet per Day
                if 'hari' in df_jadwal.columns:
                    days = df_jadwal['hari'].unique()
                    for day in days:
                        df_day = df_jadwal[df_jadwal['hari'] == day].sort_values('jam_ke')
                        df_day.to_excel(writer, index=False, sheet_name=str(day)[:31])
                
                # Sheet per Teacher
                teacher_list = sorted(list(set(df_jadwal['guru'].unique())))
                for teacher in teacher_list:
                    df_t = df_jadwal[df_jadwal['guru'] == teacher].sort_values(['hari', 'jam_ke'])
                    safe_name = str(teacher).replace(":", "").replace("/", "")[:31]
                    df_t.to_excel(writer, index=False, sheet_name=safe_name)

            st.download_button(
                label="📥 Download Excel Lengkap",
                data=buffer.getvalue(),
                file_name="Jadwal_Pelajaran.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            st.divider()
            st.subheader("👤 Jadwal Per Guru")
            
            # Get unique teachers (using the already defined list if available, or redefine)
            teacher_list = sorted(list(set(df_jadwal['guru'].unique())))
            selected_teacher = st.selectbox("Pilih Guru:", ["- Pilih -"] + teacher_list)
            
            if selected_teacher and selected_teacher != "- Pilih -":
                df_teacher = df_jadwal[df_jadwal['guru'] == selected_teacher].sort_values(['hari', 'jam_ke'])
                st.write(f"**Jadwal: {selected_teacher}**")
                st.dataframe(df_teacher, use_container_width=True)
                
                # Excel Download for Specific Teacher
                buffer_teacher = io.BytesIO()
                with pd.ExcelWriter(buffer_teacher, engine='openpyxl') as writer:
                    df_teacher.to_excel(writer, index=False, sheet_name=selected_teacher[:31])
                    df_jadwal.to_excel(writer, index=False, sheet_name="Master Jadwal")
                
                st.download_button(
                    label=f"📥 Download Jadwal {selected_teacher}",
                    data=buffer_teacher.getvalue(),
                    file_name=f"Jadwal_{selected_teacher.replace(' ', '_')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            
            # === NEW: IMPROVED CONFLICT DISPLAY ===
            if 'main_jadwal_conflicts' in st.session_state:
                conflict_data = st.session_state['main_jadwal_conflicts']
                
                # Separate hard vs soft conflicts
                hard_conflicts = conflict_data.get('hard_conflicts', [])
                soft_conflicts = conflict_data.get('soft_conflicts', [])
                
                # === HARD CONFLICTS (CRITICAL) ===
                if hard_conflicts:
                    st.error(f"🚨 **{len(hard_conflicts)} BENTROK KRITIS** (Harus Diperbaiki!)")
                    
                    with st.expander("Lihat Detail & Perbaiki", expanded=True):
                        for idx, conf in enumerate(hard_conflicts):
                            st.markdown(f"### Bentrok #{idx+1}")
                            
                            if conf['type'] == 'guru_bentrok':
                                st.error(f"""
                                **Bentrok Guru:** {conf['guru']}  
                                📅 **{conf['hari']} Jam {conf['jam_ke']}**  
                                - 🏫 Mengajar **{conf['kelas_1']}** ({conf['mapel_1']})  
                                - 🏫 Mengajar **{conf['kelas_2']}** ({conf['mapel_2']})  
                                """)
                                
                                # Smart Auto-Fix
                                render_smart_fix(
                                    conflict=conf,
                                    guru=conf['guru'],
                                    kelas_to_move=conf['kelas_2'],  # Move second class
                                    idx=idx
                                )
                            
                            elif conf['type'] == 'kelas_bentrok':
                                st.error(f"""
                                **Bentrok Kelas:** {conf['kelas']}  
                                📅 **{conf['hari']} Jam {conf['jam_ke']}**  
                                - 👨‍🏫 Diajar **{conf['guru_1']}** ({conf['mapel_1']})  
                                - 👨‍🏫 Diajar **{conf['guru_2']}** ({conf['mapel_2']})  
                                """)
                                
                                # Smart Auto-Fix for class conflict
                                render_smart_fix(
                                    conflict=conf,
                                    guru=conf['guru_2'],  # Move second teacher's session
                                    kelas_to_move=conf['kelas'],
                                    idx=idx
                                )
                            
                            st.markdown("---")
                
                # === SOFT CONFLICTS (WARNINGS) ===
            # === LOGS & STATUS ===
            with st.expander("🎓 Detail Proses & Log AI", expanded=False):
                if 'logs' in result:
                    for log in result['logs']:
                        if "Error" in log or "Failed" in log or "Critical" in log:
                            st.error(f"❌ {log}")
                        elif "Warning" in log:
                            st.warning(f"⚠️ {log}")
                        else:
                            st.info(f"🔹 {log}")

            if st.button("🗑️ Reset Jadwal"):
                del st.session_state['main_jadwal_result']
                if 'main_jadwal_conflicts' in st.session_state:
                    del st.session_state['main_jadwal_conflicts']
                st.rerun()

